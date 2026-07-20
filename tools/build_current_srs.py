from __future__ import annotations

import shutil
from datetime import date
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\School materials\SWT301\CustomerFeedbackManagementSystem")
OUT_DIR = ROOT / "docs"
IMG_DIR = ROOT / "tmp" / "srs_current_assets"
SOURCE_IMG_DIR = ROOT / "tmp" / "srs_images"
OUT_DOCX = OUT_DIR / "Customer_Feedback_Management_System_SRS_v1.4_Current_Project.docx"

BLUE = RGBColor(31, 77, 120)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
INK = RGBColor(20, 30, 45)
MUTED = RGBColor(90, 98, 110)


def ensure_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    IMG_DIR.mkdir(parents=True, exist_ok=True)


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def rounded_rect(draw: ImageDraw.ImageDraw, box, fill, outline, width=2, radius=16):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def center_text(draw: ImageDraw.ImageDraw, box, text, font, fill=(20, 30, 45), spacing=4):
    x1, y1, x2, y2 = box
    max_width = x2 - x1 - 24
    words = text.split()
    lines = []
    line = ""
    for word in words:
        test = f"{line} {word}".strip()
        if draw.textbbox((0, 0), test, font=font)[2] <= max_width:
            line = test
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    total_height = sum(draw.textbbox((0, 0), ln, font=font)[3] for ln in lines) + spacing * (len(lines) - 1)
    y = y1 + ((y2 - y1) - total_height) / 2
    for ln in lines:
        bbox = draw.textbbox((0, 0), ln, font=font)
        draw.text((x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2, y), ln, font=font, fill=fill)
        y += (bbox[3] - bbox[1]) + spacing


def arrow(draw, start, end, fill=(20, 30, 45), width=3):
    draw.line([start, end], fill=fill, width=width)
    # simple arrow head
    ex, ey = end
    sx, sy = start
    import math

    angle = math.atan2(ey - sy, ex - sx)
    for delta in (2.65, -2.65):
        x = ex + 14 * math.cos(angle + delta)
        y = ey + 14 * math.sin(angle + delta)
        draw.line([(ex, ey), (x, y)], fill=fill, width=width)


def make_architecture_diagram() -> Path:
    path = IMG_DIR / "current_architecture.png"
    img = Image.new("RGB", (1600, 1000), "white")
    d = ImageDraw.Draw(img)
    title = load_font(42, True)
    h = load_font(25, True)
    body = load_font(20)
    small = load_font(18)
    d.text((420, 30), "System Architecture - Current Implementation", font=title, fill=(13, 38, 76))

    rounded_rect(d, (40, 130, 300, 610), "#F7FAFF", "#2E74B5", 3)
    center_text(d, (40, 145, 300, 195), "Actors", h, (13, 38, 76))
    actors = ["Guest", "Customer", "Support Staff", "Department Manager", "System Admin"]
    y = 220
    for actor in actors:
        rounded_rect(d, (75, y, 265, y + 58), "#FFFFFF", "#C8D7EF", 2, 10)
        center_text(d, (75, y, 265, y + 58), actor, small)
        y += 75

    rounded_rect(d, (390, 120, 1180, 300), "#F7FAFF", "#2E74B5", 3)
    center_text(d, (410, 138, 590, 182), "Frontend", h, (13, 38, 76))
    center_text(d, (610, 135, 1135, 190), "React + Vite SPA", h, (13, 38, 76))
    for i, txt in enumerate(["Role-based routes and dashboards", "Feedback, admin, report and notification pages", "Axios client with JWT refresh handling"]):
        d.text((455, 202 + i * 30), f"- {txt}", font=body, fill=(20, 30, 45))

    rounded_rect(d, (390, 380, 1180, 690), "#F8FFFA", "#2A7A55", 3)
    center_text(d, (410, 400, 590, 450), "Backend", h, (20, 95, 65))
    center_text(d, (610, 395, 1135, 450), "ASP.NET Core 9 REST API", h, (20, 95, 65))
    boxes = [
        ("Controllers", 430, 480),
        ("Application Services", 675, 480),
        ("JWT + Google Auth", 920, 480),
        ("EF Core Repositories", 430, 580),
        ("SignalR Notifications", 675, 580),
        ("Supabase Storage", 920, 580),
    ]
    for txt, x, y in boxes:
        rounded_rect(d, (x, y, x + 205, y + 65), "#FFFFFF", "#89BFA4", 2, 10)
        center_text(d, (x, y, x + 205, y + 65), txt, small)

    rounded_rect(d, (1270, 130, 1560, 690), "#F7FAFF", "#2E74B5", 3)
    center_text(d, (1270, 145, 1560, 205), "Supabase", h, (13, 38, 76))
    for i, txt in enumerate(["PostgreSQL database", "cfms-attachments bucket", "Connection via session pooler", "Database constraints and indexes"]):
        rounded_rect(d, (1305, 235 + i * 82, 1525, 295 + i * 82), "#FFFFFF", "#C8D7EF", 2, 10)
        center_text(d, (1305, 235 + i * 82, 1525, 295 + i * 82), txt, small)

    rounded_rect(d, (390, 770, 1180, 930), "#FBFAFF", "#7B61B6", 3)
    center_text(d, (410, 792, 590, 840), "Verification", h, (80, 55, 120))
    for i, txt in enumerate(["xUnit service/model tests: 77 passing", "Vite lint and production build passing", "EF model has no pending migration changes"]):
        d.text((455, 848 + i * 28), f"- {txt}", font=body, fill=(20, 30, 45))

    arrow(d, (300, 365), (390, 210))
    arrow(d, (785, 300), (785, 380))
    arrow(d, (1180, 535), (1270, 400))
    arrow(d, (785, 690), (785, 770))
    img.save(path)
    return path


def make_data_model_diagram() -> Path:
    path = IMG_DIR / "current_data_model.png"
    img = Image.new("RGB", (1600, 1050), "white")
    d = ImageDraw.Draw(img)
    title = load_font(40, True)
    h = load_font(22, True)
    small = load_font(17)
    d.text((460, 24), "Current Data Model - Supabase PostgreSQL", font=title, fill=(13, 38, 76))

    def entity(name, fields, x, y, w=270):
        height = 48 + 25 * len(fields)
        rounded_rect(d, (x, y, x + w, y + height), "#FFFFFF", "#2E74B5", 3, 12)
        d.rectangle((x, y, x + w, y + 42), fill="#0B3D75", outline="#0B3D75")
        bbox = d.textbbox((0, 0), name, font=h)
        d.text((x + (w - (bbox[2] - bbox[0])) / 2, y + 8), name, font=h, fill="white")
        for i, field in enumerate(fields):
            d.text((x + 14, y + 54 + i * 25), field, font=small, fill=(20, 30, 45))
        return (x, y, x + w, y + height)

    e_user = entity("users", ["Id PK", "Email, password hash", "Role, Status", "DepartmentId FK", "GoogleSubject", "Soft delete fields"], 40, 120)
    e_dept = entity("departments", ["Id PK", "Name", "Description", "IsActive"], 40, 520)
    e_cat = entity("feedback_categories", ["Id PK", "Name", "Description", "DepartmentId FK", "IsActive"], 410, 120)
    e_fb = entity("feedbacks", ["Id PK", "SubmittedByUserId FK", "CategoryId FK", "DepartmentId FK", "AssignedToUserId FK", "Title, Description", "Priority, Status, Rating"], 760, 120, 310)
    e_assign = entity("feedback_assignments", ["Id PK", "FeedbackId FK", "AssignedToUserId FK", "AssignedByUserId FK", "IsActive", "Unique active per feedback"], 1190, 110, 330)
    e_resp = entity("feedback_responses", ["Id PK", "FeedbackId FK", "RespondedByUserId FK", "Content", "IsInternal, IsDeleted"], 1190, 390, 330)
    e_comment = entity("feedback_comments", ["Id PK", "FeedbackId FK", "AuthorUserId FK", "ParentCommentId FK", "Content", "IsDeleted"], 760, 500, 310)
    e_attach = entity("feedback_attachments", ["Id PK", "FeedbackId FK", "UploadedByUserId FK", "StorageKey", "ContentType, Size"], 410, 520)
    e_hist = entity("feedback_status_history", ["Id PK", "FeedbackId FK", "FromStatus, ToStatus", "ChangedByUserId FK", "Reason"], 40, 800, 330)
    e_notif = entity("notifications", ["Id PK", "UserId FK", "Type", "Title, Message", "EntityId, EntityType", "IsRead"], 570, 800, 300)
    e_audit = entity("audit_logs", ["Id PK", "UserId FK nullable", "Action", "EntityType, EntityId", "OldValues, NewValues", "IpAddress"], 1030, 800, 300)
    entity("refresh_tokens", ["Id PK", "UserId FK", "Token", "Expiry/Revoke fields", "CreatedByIp"], 40, 330, 270)

    for start, end in [
        ((310, 220), (760, 235)),
        ((680, 220), (760, 245)),
        ((310, 580), (410, 220)),
        ((310, 580), (760, 265)),
        ((1070, 250), (1190, 190)),
        ((1070, 280), (1190, 470)),
        ((920, 500), (920, 390)),
        ((550, 520), (820, 390)),
        ((370, 860), (760, 330)),
        ((570, 875), (310, 240)),
        ((1030, 875), (310, 250)),
    ]:
        arrow(d, start, end, fill=(75, 85, 100), width=2)

    d.text((1100, 1010), "Note: diagram is conceptual; exact columns are defined by EF Core migrations.", font=small, fill=(80, 80, 80))
    img.save(path)
    return path


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(9.5)
    r.font.name = "Calibri"
    if color:
        r.font.color.rgb = color


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_text(cell, header, True, BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_image(doc: Document, image_path: Path, caption: str, width: float = 6.4) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    inline_shape = run.add_picture(str(image_path), width=Inches(width))
    inline_shape._inline.docPr.set("title", caption)
    inline_shape._inline.docPr.set("descr", caption)
    add_caption(doc, caption)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def set_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = "Customer Feedback Management System SRS"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header_p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.text = "Prepared for SWT301 project documentation"
    for run in footer_p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


def make_doc() -> None:
    ensure_dirs()
    arch = make_architecture_diagram()
    data_model = make_data_model_diagram()
    copied_images = {}
    for image_name in ["image1.png", "image4.png", "image5.png", "image6.png", "image7.png"]:
        src = SOURCE_IMG_DIR / image_name
        dst = IMG_DIR / image_name
        if src.exists():
            shutil.copyfile(src, dst)
            copied_images[image_name] = dst

    doc = Document()
    configure_styles(doc)
    set_header_footer(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(48)
    title.paragraph_format.space_after = Pt(6)
    r = title.add_run("SOFTWARE REQUIREMENTS SPECIFICATION")
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor(13, 38, 76)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Customer Feedback Management System")
    r.font.size = Pt(18)
    r.font.bold = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Version 1.4 | Current Project Baseline | Generated from implemented BE/FE repository").italic = True
    doc.add_paragraph()
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Project type", "Web application with React frontend and ASP.NET Core REST API"],
            ["Backend", ".NET 9, ASP.NET Core, EF Core, JWT Bearer authentication, SignalR"],
            ["Database and storage", "Supabase PostgreSQL, Supabase Storage bucket cfms-attachments"],
            ["Verification status", "Backend build passed, 77 xUnit tests passed, frontend lint/build passed"],
            ["Date", date.today().isoformat()],
        ],
        [1.8, 4.6],
    )
    doc.add_page_break()

    doc.add_heading("Revision History", level=1)
    add_table(
        doc,
        ["Version", "Date", "Main Changes", "Author/Team"],
        [
            ["1.0", "Initial", "Initial SRS draft for Customer Feedback System.", "Project Team"],
            ["1.1", "Updated", "Added user roles, permission matrix, workflows, requirements, database planning, and tests.", "Project Team"],
            ["1.2", "Updated", "Added authentication controls, attachment constraints, reporting, and consistency fixes.", "Project Team"],
            ["1.3", "Updated", "Clarified MVP scope and kept core management/reporting features.", "Project Team"],
            ["1.4", date.today().isoformat(), "Rebuilt to match the current implemented React, ASP.NET Core, EF Core, and Supabase project.", "Codex + Project Team"],
        ],
        [0.75, 0.95, 3.85, 0.95],
    )

    doc.add_heading("Chapter 1: Introduction", level=1)
    doc.add_heading("1.1 Project Overview", level=2)
    doc.add_paragraph(
        "The Customer Feedback Management System (CFMS) is a web application that centralizes customer feedback intake, assignment, handling, response, notification, reporting, and administrative governance. The implemented system supports a React/Vite frontend, an ASP.NET Core 9 REST API, EF Core persistence, Supabase PostgreSQL, Supabase Storage, JWT authentication, Google login, and SignalR-based real-time notification delivery."
    )
    doc.add_paragraph(
        "The system replaces scattered feedback channels with a controlled ticket workflow. Customers submit and track their own feedback; support staff handle assigned tickets; department managers supervise department work; and system administrators manage users, categories, departments, reports, and audit logs."
    )
    doc.add_heading("1.2 Purpose", level=2)
    doc.add_paragraph(
        "This SRS defines the implemented project baseline and the requirements that the current source code satisfies. It is intended to support grading, testing, maintenance, and future enhancement decisions."
    )
    doc.add_heading("1.3 Scope", level=2)
    add_bullets(
        doc,
        [
            "Customer account registration, login, logout, Google login, profile update, and password change.",
            "Role-based dashboards and navigation for Customer, Support Staff, Department Manager, and System Admin.",
            "Customer feedback submission with required category and optional attachments.",
            "Feedback workflow: assignment, reassignment, unassignment, status changes, priority changes, responses, comments, attachments, and rating.",
            "Department and feedback category management.",
            "User management with role/status changes, soft deletion, and token revocation.",
            "In-app and real-time notifications, audit logs, summary reports, and staff workload reports.",
            "Database integrity rules enforced through EF Core migrations and PostgreSQL constraints.",
        ],
    )
    doc.add_heading("1.4 Out of Scope / Future Enhancements", level=2)
    add_bullets(
        doc,
        [
            "Email notification delivery beyond current in-app and SignalR notification flow.",
            "Forgot-password email flow and two-factor authentication.",
            "CSV/PDF report export.",
            "AI sentiment analysis, chatbot support, social media integration, and native mobile apps.",
            "Payment, refund, CRM, or external ticketing integrations.",
        ],
    )
    doc.add_heading("1.5 Definitions", level=2)
    add_table(
        doc,
        ["Term", "Meaning"],
        [
            ["Feedback", "A customer-submitted ticket containing title, description, category, status, priority, optional attachments, responses, comments, and rating."],
            ["Customer", "Registered external user who submits and tracks their own feedback."],
            ["Support Staff", "Internal user who can handle only feedback assigned to them."],
            ["Department Manager", "Internal supervisor who can manage and report on feedback within their department."],
            ["System Admin", "Platform-level administrator with global management and audit capabilities."],
            ["Supabase", "Hosted PostgreSQL and Storage platform used by the current implementation."],
        ],
        [1.35, 5.05],
    )

    doc.add_heading("Chapter 2: Overall Description", level=1)
    doc.add_heading("2.1 User Classes and Permissions", level=2)
    add_table(
        doc,
        ["Function", "Guest", "Customer", "Support Staff", "Dept. Manager", "System Admin"],
        [
            ["Register account", "Yes", "No", "No", "No", "Can create/manage users"],
            ["Login / logout / refresh token", "Login only", "Yes", "Yes", "Yes", "Yes"],
            ["Google login", "Yes", "Yes", "Yes", "Yes", "Yes"],
            ["Submit feedback", "No", "Yes", "No", "No", "No"],
            ["View feedback", "No", "Own only", "Assigned only", "Department only", "All"],
            ["Respond/comment", "No", "Own comments", "Assigned tickets", "Department tickets", "All/moderation"],
            ["Assign/reassign/unassign", "No", "No", "No", "Department only", "All"],
            ["Change status/priority", "No", "No", "Assigned only", "Department only", "All"],
            ["Rate feedback", "No", "Own resolved/closed only", "No", "No", "No"],
            ["Manage users", "No", "Profile only", "Profile only", "Profile only", "Yes"],
            ["Manage departments/categories", "No", "Read active categories", "Read active categories", "Read departments", "Yes"],
            ["Reports", "No", "No", "No", "Department scoped", "Global"],
            ["Audit logs", "No", "No", "No", "No", "Yes"],
        ],
        [1.65, 0.75, 0.95, 1.05, 1.05, 1.1],
    )
    if "image1.png" in copied_images:
        add_image(doc, copied_images["image1.png"], "Figure 1. Original high-level use-case diagram retained for actor context.", 6.5)

    doc.add_heading("2.2 Operating Environment", level=2)
    add_table(
        doc,
        ["Layer", "Current Implementation"],
        [
            ["Frontend", "React 19, Vite 8, React Router 7, Axios, Bootstrap styling, role-based route guards."],
            ["Backend", "ASP.NET Core 9 Web API, FluentValidation, AutoMapper, JWT Bearer authentication, SignalR hub."],
            ["Application", "Service-layer business rules for auth, users, feedback, assignments, comments, responses, notifications, reports, departments, and categories."],
            ["Persistence", "EF Core 9 with Npgsql against Supabase PostgreSQL; migrations define constraints and indexes."],
            ["Storage", "Supabase Storage bucket cfms-attachments; server-side upload/delete via configured secret key."],
            ["Testing", "xUnit service/model tests and frontend lint/build verification."],
        ],
        [1.45, 4.95],
    )
    add_image(doc, arch, "Figure 2. Current implementation architecture.", 6.5)
    doc.add_heading("2.3 Design and Implementation Constraints", level=2)
    add_bullets(
        doc,
        [
            "All protected backend endpoints require authentication and server-side role validation.",
            "JWT access tokens are rejected after the user is disabled or deleted.",
            "Passwords are stored as hashes only.",
            "Customers can only access feedback that they submitted.",
            "Support Staff can only handle assigned feedback.",
            "Department Managers are restricted to feedback and staff in their department.",
            "System Admin cannot disable, demote, or delete their own active admin account.",
            "Support Staff with active feedback cannot be disabled, deleted, or moved until tickets are reassigned or unassigned.",
        ],
    )

    doc.add_heading("Chapter 3: Functional Requirements", level=1)
    doc.add_heading("3.1 Functional Requirement Catalogue", level=2)
    add_table(
        doc,
        ["ID", "Requirement"],
        [
            ["FR-AUTH-01", "The system shall allow customer registration with email, name, phone, and password validation."],
            ["FR-AUTH-02", "The system shall support email/password login, Google login, JWT issuing, refresh tokens, logout, and password change."],
            ["FR-AUTH-03", "The system shall reject disabled or deleted accounts for login and protected API use."],
            ["FR-USER-01", "The system shall allow authenticated users to view and update their own profile."],
            ["FR-USER-02", "The system shall allow System Admin to list, view, update role, update status, deactivate, reactivate, and soft-delete users."],
            ["FR-FB-01", "The system shall allow Customers to submit feedback with title, description, required active category, and optional validated attachments."],
            ["FR-FB-02", "The system shall initialize new feedback with status New and priority Medium."],
            ["FR-FB-03", "The system shall allow scoped feedback listing and detail viewing according to actor role."],
            ["FR-FB-04", "The system shall allow internal users to update feedback content, category, and priority according to scope and ticket state."],
            ["FR-FB-05", "The system shall enforce the feedback status lifecycle and require reasons for Rejected and Closed."],
            ["FR-FB-06", "The system shall allow Customers to rate their own Resolved or Closed feedback from 1 to 5."],
            ["FR-ASSIGN-01", "The system shall allow Department Manager and System Admin to assign, reassign, unassign, and view assignment history under scope rules."],
            ["FR-COMM-01", "The system shall allow scoped comments and official staff responses, with internal responses hidden from Customers."],
            ["FR-ATT-01", "The system shall allow upload and deletion of validated attachments while feedback is not Closed or Rejected."],
            ["FR-CAT-01", "The system shall allow active category lookup and System Admin category CRUD/update operations."],
            ["FR-DEPT-01", "The system shall allow Department Manager/System Admin active department lookup and System Admin department management."],
            ["FR-NOTIF-01", "The system shall create, list, mark-read, and mark-all-read notifications for the authenticated user."],
            ["FR-REPORT-01", "The system shall provide summary, status, category, priority, trend, and staff workload reports for Manager/Admin roles."],
            ["FR-AUDIT-01", "The system shall log important create, update, delete, assignment, status, role, and account actions."],
        ],
        [1.05, 5.35],
    )
    doc.add_heading("3.2 Key User Stories", level=2)
    stories = [
        ("US-01 Customer Registration", "As a Guest, I want to create a Customer account so that I can submit and track feedback."),
        ("US-02 Login and Session Management", "As a user, I want secure login/logout and token refresh so that my session remains controlled."),
        ("US-03 Submit Feedback", "As a Customer, I want to submit feedback with a category and optional attachment so that the organization can process my request."),
        ("US-04 Track Feedback", "As a Customer, I want to view my own feedback, comments, responses, status history, attachments, and rating state."),
        ("US-05 Handle Assigned Feedback", "As Support Staff, I want to respond to assigned feedback and update workflow status."),
        ("US-06 Manage Department Feedback", "As Department Manager, I want to assign work and monitor feedback in my department."),
        ("US-07 Administer the Platform", "As System Admin, I want to manage users, categories, departments, audit logs, and global reports."),
        ("US-08 Receive Notifications", "As a user, I want to see notifications when feedback activity requires my attention."),
    ]
    for heading, body in stories:
        doc.add_heading(heading, level=3)
        doc.add_paragraph(body)

    doc.add_heading("3.3 Workflow Rules", level=2)
    if "image4.png" in copied_images:
        add_image(doc, copied_images["image4.png"], "Figure 3. Feedback status lifecycle retained from the original SRS.", 6.5)
    add_table(
        doc,
        ["Status", "Meaning", "Allowed Next Status"],
        [
            ["New", "Feedback submitted and not yet assigned.", "Assigned, Rejected"],
            ["Assigned", "Feedback assigned to support staff.", "InProgress, Rejected"],
            ["InProgress", "Staff is actively handling feedback.", "WaitingForCustomer, Resolved, Rejected"],
            ["WaitingForCustomer", "Staff requires customer information.", "InProgress, Resolved, Closed"],
            ["Resolved", "Issue resolved and awaiting confirmation.", "Closed, InProgress"],
            ["Rejected", "Feedback rejected with a reason.", "Closed"],
            ["Closed", "Feedback process is complete.", "None"],
        ],
        [1.35, 2.45, 2.6],
    )
    if "image5.png" in copied_images:
        add_image(doc, copied_images["image5.png"], "Figure 4. Customer feedback submission activity diagram retained.", 6.2)
    if "image6.png" in copied_images:
        add_image(doc, copied_images["image6.png"], "Figure 5. Feedback handling activity diagram retained.", 6.2)
    if "image7.png" in copied_images:
        add_image(doc, copied_images["image7.png"], "Figure 6. Admin management activity diagram retained with expanded requirements text.", 6.2)

    doc.add_heading("Chapter 4: External Interface Requirements", level=1)
    doc.add_heading("4.1 Frontend Pages", level=2)
    add_table(
        doc,
        ["Page", "Primary Actors", "Purpose"],
        [
            ["Login / Register", "Guest", "Authenticate or create a customer account."],
            ["Dashboard", "All authenticated roles", "Show role-specific summary and navigation."],
            ["Submit Feedback", "Customer", "Create feedback with active category and optional attachments."],
            ["My Feedbacks", "Customer", "List and filter own feedback."],
            ["Assigned Feedbacks", "Support Staff, Department Manager, System Admin", "List scoped internal feedback work."],
            ["Feedback Detail", "Scoped actors", "View and act on feedback details, responses, comments, assignments, status, priority, attachments, and rating."],
            ["Reports", "Department Manager, System Admin", "View summary, trend, and staff workload analytics."],
            ["Manage Users", "System Admin", "Administer user roles, status, departments, and soft deletion."],
            ["Manage Categories / Departments", "System Admin", "Maintain operational configuration."],
            ["Notifications / Profile / Audit Logs", "Scoped roles", "Notification handling, own profile, and admin audit review."],
        ],
        [1.65, 1.85, 2.9],
    )
    doc.add_heading("4.2 REST API Summary", level=2)
    add_table(
        doc,
        ["Area", "Representative Endpoints"],
        [
            ["Authentication", "POST /api/auth/register, /login, /google-login, /refresh-token, /logout, /change-password"],
            ["Users", "GET /api/users/me, PUT /api/users/me, GET /api/admin/users, PATCH /api/admin/users/{id}/role, PATCH /api/users/{id}/status"],
            ["Feedback", "GET /api/feedback, GET /api/feedback/my, POST /api/feedback, PUT/PATCH/DELETE /api/feedback/{id}, rating and attachment routes"],
            ["Assignments", "PATCH /api/feedback/{feedbackId}/assign, /reassign, GET/DELETE /api/feedback/{feedbackId}/assignments"],
            ["Comments and responses", "GET/POST/PUT/DELETE /api/feedback/{feedbackId}/comments and /responses"],
            ["Configuration", "GET /api/categories, /api/departments, and /api/admin/categories or /api/admin/departments management routes"],
            ["Reports and audit", "GET /api/reports/*, GET /api/admin/audit-logs"],
            ["Notifications", "GET /api/notifications, unread count, mark read, mark all read"],
        ],
        [1.55, 4.85],
    )

    doc.add_heading("Chapter 5: Data Requirements", level=1)
    add_image(doc, data_model, "Figure 7. Current conceptual data model based on EF Core entities and migrations.", 6.5)
    doc.add_heading("5.1 Main Tables", level=2)
    add_table(
        doc,
        ["Table", "Purpose"],
        [
            ["users", "Stores account identity, role, status, department, Google subject, password hash, and soft-delete metadata."],
            ["refresh_tokens", "Stores refresh token lifecycle for session management and revocation."],
            ["departments", "Stores operational departments used for staff and feedback scoping."],
            ["feedback_categories", "Stores active/inactive categories and optional department ownership."],
            ["feedbacks", "Stores feedback core data, ownership, assignment pointer, department, status, priority, rating, and timestamps."],
            ["feedback_assignments", "Stores assignment history and enforces one active assignment per feedback."],
            ["feedback_responses", "Stores official staff responses, including internal responses hidden from Customers."],
            ["feedback_comments", "Stores discussion comments and threaded replies."],
            ["feedback_attachments", "Stores attachment metadata and Supabase Storage keys."],
            ["feedback_status_history", "Stores immutable status transition history."],
            ["notifications", "Stores in-app notification records and read state."],
            ["audit_logs", "Stores important business action audit data."],
        ],
        [1.75, 4.65],
    )
    doc.add_heading("5.2 Database Integrity Requirements", level=2)
    add_bullets(
        doc,
        [
            "Feedback CategoryId is required.",
            "Feedback rating must be null or between 1 and 5.",
            "Feedback status and priority must be one of the defined enum values.",
            "User role, user status, notification type, and status history values must be constrained to valid enums.",
            "Only one active feedback assignment is allowed per feedback.",
            "Soft-deleted feedback is excluded from normal feedback-related queries.",
        ],
    )
    doc.add_heading("5.3 Attachment Rules", level=2)
    add_bullets(
        doc,
        [
            "Maximum three attachments per feedback.",
            "Maximum file size is 5 MB.",
            "Allowed extensions: jpg, jpeg, png, gif, pdf, docx, xlsx.",
            "Content type must match the file extension.",
            "Attachment file name length must be between 1 and 256 characters.",
            "Attachments cannot be added to or deleted from Closed or Rejected feedback except by System Admin where deletion is allowed by service rule.",
        ],
    )

    doc.add_heading("Chapter 6: Non-Functional Requirements", level=1)
    add_table(
        doc,
        ["ID", "Requirement"],
        [
            ["NFR-SEC-01", "All protected APIs shall require JWT authentication."],
            ["NFR-SEC-02", "Backend services shall enforce authorization and ownership; UI route guards are not sufficient."],
            ["NFR-SEC-03", "Disabled or deleted users shall not be allowed to continue using protected APIs with old access tokens."],
            ["NFR-SEC-04", "Passwords and JWT secrets shall not be stored in source code; development secrets use .NET user-secrets."],
            ["NFR-REL-01", "Feedback submission and workflow actions shall persist before success is returned."],
            ["NFR-REL-02", "Notification failures shall not block the core feedback transaction."],
            ["NFR-MAIN-01", "The solution shall keep Domain, Application, Infrastructure, API, and frontend layers separated."],
            ["NFR-MAIN-02", "Business rules shall be testable in application services."],
            ["NFR-USE-01", "Forms shall show clear validation errors and protect users from invalid workflow actions."],
            ["NFR-AUDIT-01", "Important business mutations shall write audit log records where supported."],
        ],
        [1.1, 5.3],
    )

    doc.add_heading("Chapter 7: Verification and Acceptance", level=1)
    doc.add_heading("7.1 Acceptance Test Baseline", level=2)
    add_table(
        doc,
        ["Verification Item", "Current Result"],
        [
            ["Backend Release build", "Passed with 0 warnings and 0 errors."],
            ["Backend xUnit tests", "77 passed, 0 failed, 0 skipped."],
            ["EF Core migration drift", "No pending model changes."],
            ["Frontend ESLint", "Passed."],
            ["Frontend production build", "Passed."],
            [".NET vulnerable package scan", "No vulnerable packages reported from configured NuGet sources."],
            ["npm audit", "Not completed in Codex environment because advisory scan exports dependency metadata to npm external service."],
        ],
        [2.2, 4.2],
    )
    doc.add_heading("7.2 Actor Acceptance Checklist", level=2)
    add_table(
        doc,
        ["Actor", "Acceptance Conditions"],
        [
            ["Customer", "Can register/login, submit feedback, view only own feedback, comment, upload/delete allowed own attachments, rate resolved/closed feedback, and delete own New feedback."],
            ["Support Staff", "Can view and handle only assigned feedback, respond, comment, update status, update priority, and cannot manage unassigned tickets."],
            ["Department Manager", "Can assign/reassign/unassign and report only inside own department; cannot manage global users or audit logs."],
            ["System Admin", "Can manage users, categories, departments, audit logs, reports, and all feedback; cannot disable/demote/delete own active admin account."],
        ],
        [1.45, 4.95],
    )

    doc.add_heading("Chapter 8: Traceability", level=1)
    add_table(
        doc,
        ["Feature Area", "Implemented In"],
        [
            ["Authentication", "AuthController, AuthService, JwtService, GoogleAuthService, RefreshTokenRepository"],
            ["Feedback workflow", "FeedbackController, FeedbackService, FeedbackStatusRules, FeedbackRepository"],
            ["Assignments", "AssignmentsController, FeedbackAssignmentService, FeedbackAssignment entity/configuration"],
            ["Comments and responses", "CommentsController, ResponsesController, FeedbackCommentService, FeedbackResponseService"],
            ["Categories and departments", "CategoriesController, DepartmentsController, FeedbackCategoryService, DepartmentService"],
            ["Reports", "ReportsController, ReportService"],
            ["Notifications", "NotificationsController, NotificationService, SignalR NotificationHub"],
            ["Audit logs", "AuditLogsController, AuditLogService, AuditLogRepository"],
            ["Frontend UI", "React pages under FE/src/pages and route guards under FE/src/routes"],
            ["Persistence", "AppDbContext, EF Core configurations, migrations"],
            ["Tests", "BE/tests/CFMS.Tests service, model, storage, category, department, notification, audit, and actor-scope tests"],
        ],
        [1.75, 4.65],
    )

    doc.add_heading("Chapter 9: Appendix", level=1)
    doc.add_heading("9.1 Notes on Reused Diagrams", level=2)
    add_bullets(
        doc,
        [
            "The original use-case diagram is retained as a high-level actor overview, but the written requirements are authoritative for newer functions.",
            "The original status lifecycle and activity diagrams are retained because they still match the implemented workflow at a business level.",
            "The original architecture and ERD diagrams were replaced because they referenced SQL Server/Spring Boot-era design and did not match the current Supabase/.NET implementation.",
        ],
    )
    doc.add_heading("9.2 Current Risk Register", level=2)
    add_table(
        doc,
        ["Risk", "Status / Mitigation"],
        [
            ["External npm advisory audit not executed", "Requires explicit approval to send dependency metadata to npm. Local package tree, lint, and build are clean."],
            ["Email notification channel", "Out of scope; in-app and SignalR notifications are implemented."],
            ["Forgot-password workflow", "Out of scope for current baseline."],
            ["Large-data performance", "Pagination exists; load testing with production-sized data remains a future activity."],
        ],
        [2.1, 4.3],
    )

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    make_doc()
