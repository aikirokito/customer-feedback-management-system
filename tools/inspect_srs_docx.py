import sys
from pathlib import Path

from docx import Document

sys.stdout.reconfigure(encoding="utf-8")

source = Path(r"D:\School materials\SWT301\Customer_Feedback_Management_System_SRS_v1.3.docx")
doc = Document(source)

print("paragraphs", len(doc.paragraphs), "tables", len(doc.tables), "inline_shapes", len(doc.inline_shapes), "sections", len(doc.sections))
print("--- first paragraphs ---")
for index, paragraph in enumerate(doc.paragraphs[:140]):
    text = paragraph.text.strip()
    if text:
        print(index, paragraph.style.name, repr(text[:240]))

print("--- tables ---")
for table_index, table in enumerate(doc.tables):
    preview_rows = []
    for row in table.rows[:5]:
        preview_rows.append([cell.text.strip().replace("\n", " | ")[:100] for cell in row.cells[:6]])
    print("TABLE", table_index, "rows", len(table.rows), "cols", len(table.columns), preview_rows)
